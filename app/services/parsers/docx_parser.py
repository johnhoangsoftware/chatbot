"""
DOCX parser - wraps existing TechnicalDocumentParser for DOCX files.
"""

import os
from typing import List
from .base import BaseParser, ParsedDocument, Section, TableData, ImageData

try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

# Import the existing TechnicalDocumentParser
import sys
sys.path.insert(0, os.path.dirname(os.path.dirname(__file__)))

class DOCXParser(BaseParser):
    """Parser for DOCX files using the existing TechnicalDocumentParser."""
    
    def __init__(self):
        pass
    
    @property
    def supported_extensions(self) -> List[str]:
        return ['.docx', '.doc']
    
    def parse(self, file_path: str) -> ParsedDocument:
        """Parse a DOCX file."""
        if not self.validate(file_path):
            raise ValueError(f"Invalid or unsupported file: {file_path}")
        
        # Delegate to existing parser
        return self._parse_docx(file_path)
    
    def _parse_docx(self, file_path: str) -> ParsedDocument:
        """Parse DOCX file and extract content."""
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx library is required for DOCX parsing. Install with: pip install python-docx")
        
        doc = Document(file_path)
        
        full_content = []
        pages = []  # DOCX doesn't have pages, so we'll treat each section as a page
        all_sections = []
        all_tables = []
        all_images = []
        
        # Extract paragraphs
        current_page = []
        page_num = 1
        word_count = 0
        
        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue
            
            # Detect headings
            if para.style.name.startswith('Heading'):
                try:
                    level = int(para.style.name.replace('Heading', '').strip() or '1')
                except:
                    level = 1
                
                all_sections.append(Section(
                    title=text,
                    level=level,
                    content="",
                    page_number=page_num
                ))
            
            current_page.append(text)
            word_count += len(text.split())
            
            # Break into "pages" every ~500 words for chunking consistency
            if word_count > 500:
                page_content = "\n\n".join(current_page)
                pages.append({
                    "page_number": page_num,
                    "content": page_content,
                    "word_count": len(page_content.split()),
                    "char_count": len(page_content),
                    "has_tables": False,
                    "has_figures": False,
                    "sections": []
                })
                full_content.append(f"[Section {page_num}]\n{page_content}")
                current_page = []
                word_count = 0
                page_num += 1
        
        # Add remaining content
        if current_page:
            page_content = "\n\n".join(current_page)
            pages.append({
                "page_number": page_num,
                "content": page_content,
                "word_count": len(page_content.split()),
                "char_count": len(page_content),
                "has_tables": False,
                "has_figures": False,
                "sections": []
            })
            full_content.append(f"[Section {page_num}]\n{page_content}")
        
        # Extract tables
        for table_idx, table in enumerate(doc.tables):
            rows = []
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                rows.append(" | ".join(cells))
            
            if rows:
                table_content = "\n".join(rows)
                all_tables.append(TableData(
                    page_number=len(pages),
                    content=table_content,
                    rows=len(rows),
                    columns=len(table.columns)
                ))
                full_content.append(f"\n[Table {table_idx + 1}]\n{table_content}\n[/Table]\n")
        
        # Extract images from document
        all_images = self._extract_images_from_docx(doc)
        
        # Add image descriptions to content for better chunking context
        for idx, img in enumerate(all_images):
            full_content.append(f"\n[Image {idx + 1}: {img.description}]\n")
        
        # Build metadata
        core_props = doc.core_properties
        metadata = {
            "filename": os.path.basename(file_path),
            "file_path": file_path,
            "page_count": len(pages),
            "total_words": sum(p["word_count"] for p in pages),
            "total_chars": sum(p["char_count"] for p in pages),
            "table_count": len(all_tables),
            "section_count": len(all_sections),
            "has_toc": False,
            "docx_metadata": {
                "author": core_props.author or "",
                "title": core_props.title or "",
                "subject": core_props.subject or "",
                "created": str(core_props.created) if core_props.created else "",
                "modified": str(core_props.modified) if core_props.modified else "",
            },
            "document_type": "Word Document"
        }
        
        # Update metadata with image count
        metadata["image_count"] = len(all_images)
        
        return ParsedDocument(
            filename=os.path.basename(file_path),
            content="\n\n".join(full_content),
            pages=pages,
            metadata=metadata,
            sections=all_sections,
            tables=all_tables,
            toc=[],
            images=all_images
        )
    
    def _extract_images_from_docx(self, doc) -> List[ImageData]:
        """
        Extract all images from DOCX file.
        
        DOCX images are stored as relationships in the document.
        """
        images = []
        
        try:
            from ..image_processor import get_image_processor
            from ...config import get_settings
            
            settings = get_settings()
            processor = get_image_processor(enabled=settings.enable_image_processing)
            
            # Get all image relationships
            for rel_id, rel in doc.part.rels.items():
                if "image" in rel.reltype:
                    try:
                        # Get image bytes
                        image_bytes = rel.target_part.blob
                        
                        if processor:
                            # Analyze with VLLM - check if it's a UI design or technical diagram
                            description = processor.analyze_image(image_bytes, context="technical")
                            image_type = processor.classify_image_type(description)
                        else:
                            description = "Image (processing disabled)"
                            image_type = "figure"
                        
                        images.append(ImageData(
                            page_number=1,  # DOCX doesn't have page numbers in same way
                            image_bytes=image_bytes,
                            description=description,
                            image_type=image_type,
                            bbox=None
                        ))
                    except Exception as e:
                        import logging
                        logger = logging.getLogger(__name__)
                        logger.warning(f"Failed to extract image from DOCX: {e}")
                        continue
        
        except Exception as e:
            import logging
            logger = logging.getLogger(__name__)
            logger.error(f"Error extracting images from DOCX: {e}")
        
        return images

